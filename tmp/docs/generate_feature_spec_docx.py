from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("output/doc/iReside_System_Feature_Specification.docx")


INTRO = (
    "This document provides a codebase-grounded overview of the currently "
    "implemented features and technical capabilities in the iReside platform. "
    "It reflects the active system scope only."
)


SECTIONS = [
    (
        "I. Authentication & Role-Based Access",
        [
            (
                "1. Unified Authentication System",
                [
                    ("Account Access", "Supports login, signup, auth callback handling, and logout flows."),
                    ("Role Routing", "Automatically routes users into tenant, landlord, or admin workspaces based on account role."),
                    ("Protected Sessions", "Uses Supabase-backed session handling across app layouts and API routes."),
                ],
            ),
            (
                "2. Role-Based Platform Segmentation",
                [
                    ("Tenant Workspace", "Dedicated pages for dashboard, community, lease, payments, messages, profile, and settings."),
                    ("Landlord Workspace", "Dedicated pages for dashboard, properties, tenants, applications, invoices, maintenance, statistics, messages, profile, and settings."),
                    ("Admin Workspace", "Separate administrative surfaces for dashboard, registrations, and user management."),
                ],
            ),
        ],
    ),
    (
        "II. AI & Intelligence Engine",
        [
            (
                "1. iRis AI Assistant (RAG-Powered)",
                [
                    ("Contextual Intelligence", "Retrieves tenant-specific lease, property, unit, maintenance, and payment context before responding."),
                    ("Conversational Support", "Provides tenant-facing AI chat through the iRis chat API and tenant chat interfaces."),
                    ("History Retrieval", "Supports fetching previous iRis conversation history."),
                ],
            ),
            (
                "2. AI Safety & Assistance Tools",
                [
                    ("Message Redaction", "Uses AI-assisted redaction for sensitive message content before sending."),
                    ("Operational Guidance", "Provides landlord-facing system advisory and portfolio insight support."),
                ],
            ),
        ],
    ),
    (
        "III. Landlord Property & Unit Management",
        [
            (
                "1. Property Creation and Editing",
                [
                    ("Property Wizard", "Supports structured property setup with staged inputs for details, units, media, and review."),
                    ("Amenity Management", "Supports both predefined and custom amenities per property."),
                    ("Media Handling", "Supports upload, ordering, preview, and removal of property media."),
                ],
            ),
            (
                "2. Visual Unit Planning",
                [
                    ("Blueprint Builder", "Includes drag/drop floor planning with snapping, overlap detection, zoom, fit, and undo behavior."),
                    ("Multi-Floor Management", "Supports multiple floors, floor switching, corridor placement, and spatial editing."),
                    ("Unit Configuration", "Opens a dedicated unit wizard for configuring individual rentable units."),
                ],
            ),
            (
                "3. Smart Contract Template Builder",
                [
                    ("Clause-Based Authoring", "Lets landlords assemble lease templates through guided contract building."),
                    ("Custom Clause Support", "Supports adding and editing custom clauses."),
                    ("Template Persistence", "Stores contract templates as part of property configuration."),
                ],
            ),
        ],
    ),
    (
        "IV. Applications, Onboarding & Tenant Acquisition",
        [
            (
                "1. Tenant Application Flows",
                [
                    ("Application Submission", "Tenants can submit rental applications through structured application pages."),
                    ("Application Tracking", "Tenant interfaces reflect progress and status of submitted applications."),
                    ("Landlord Review", "Landlords can retrieve and review applications through dedicated APIs and dashboards."),
                ],
            ),
            (
                "2. Walk-In Application Workflow",
                [
                    ("In-Person Intake", "Supports landlord-assisted walk-in application capture."),
                    ("Requirements Capture", "Records applicant identity, employment, checklist, payment, and lease-related details."),
                    ("Embedded Signing Tools", "Includes signature capture, signing mode selection, and payment record entry."),
                ],
            ),
            (
                "3. Landlord Registration",
                [
                    ("Registration Flow", "Includes a dedicated become-a-landlord workflow."),
                    ("Admin Approval", "Admin can review and update landlord registration decisions."),
                ],
            ),
        ],
    ),
    (
        "V. Lease Management & Digital Signing",
        [
            (
                "1. Lease Lifecycle Management",
                [
                    ("Lease Hub", "Tenant lease area shows active lease details, documents, notices, and renewal-related states."),
                    ("Status Controls", "Uses explicit lease transition rules and validation."),
                    ("Lease Retrieval", "Dedicated APIs provide tenant lease data and summaries."),
                ],
            ),
            (
                "2. Digital Lease Signing",
                [
                    ("Signature Workflow", "Tenants can review and sign lease agreements through dedicated signing pages."),
                    ("Secure Signing", "Uses JWT-based signing support and signature validation logic."),
                    ("Completion Handling", "Supports successful signing and landlord-side lease finalization."),
                ],
            ),
            (
                "3. Lease Documentation & Auditability",
                [
                    ("Lease Rendering", "Includes reusable lease document, header, and signature components."),
                    ("Audit Trail Support", "Landlord interfaces include lease audit-trail visibility and status badges."),
                    ("Contract Preview", "Landlords can preview contracts before finalization."),
                ],
            ),
        ],
    ),
    (
        "VI. Payments, Billing & Invoicing",
        [
            (
                "1. Tenant Financial Workspace",
                [
                    ("Payment Dashboard", "Displays balances, history, overdue items, and utility-related breakdowns."),
                    ("Checkout Flow", "Includes a dedicated payment checkout experience with method selection and confirmation states."),
                    ("Lease-Linked Billing", "Connects payment information to tenant lease data."),
                ],
            ),
            (
                "2. Landlord Billing Operations",
                [
                    ("Invoice Management", "Landlords can load, filter, and inspect invoices through dashboard tools and APIs."),
                    ("Payment Overview", "Surfaces paid, pending, and overdue payment categories."),
                    ("Payment Recording", "Supports in-person payment and payment record workflows."),
                ],
            ),
            (
                "3. Financial Messaging Support",
                [
                    ("Invoice System Messages", "Messaging interfaces can render invoice-style system messages."),
                    ("Conversation Payment Context", "Message flows can retrieve payment history relevant to a conversation."),
                ],
            ),
        ],
    ),
    (
        "VII. Maintenance Operations",
        [
            (
                "1. Maintenance Request Management",
                [
                    ("Maintenance Dashboard", "Landlords can monitor requests, priorities, statuses, and summary metrics."),
                    ("Structured Request States", "Maintenance data is normalized into clear priority and status labels."),
                    ("Request Inspection", "Includes modal-based request review and management flows."),
                ],
            ),
            (
                "2. Tenant Service Context",
                [
                    ("Maintenance Awareness in AI", "Tenant maintenance data is available to the iRis assistant for contextual support."),
                    ("Cross-Workflow Linking", "Maintenance actions connect with other landlord and communication surfaces."),
                ],
            ),
        ],
    ),
    (
        "VIII. Messaging, Communication & Moderation",
        [
            (
                "1. Direct Messaging System",
                [
                    ("Conversation Management", "Supports conversation lists, unread counts, and direct conversation creation."),
                    ("Message Exchange", "Supports loading and sending messages within conversations."),
                    ("User Search", "Supports finding users and starting new direct conversations."),
                ],
            ),
            (
                "2. Rich Media & File Exchange",
                [
                    ("Attachment Uploads", "Supports file and image uploads with metadata handling."),
                    ("Drag-and-Drop Composer", "Messaging interfaces support drag/drop attachment workflows."),
                    ("Downloadable Assets", "Supports downloading generated or shared message artifacts."),
                ],
            ),
            (
                "3. Communication Safety",
                [
                    ("Sensitive Content Redaction", "Messages can be sanitized before delivery."),
                    ("User Reporting", "Messaging includes report-user flows."),
                    ("Action-State Awareness", "Conversation summaries include moderation/action-related metadata."),
                ],
            ),
        ],
    ),
    (
        "IX. Community Hub",
        [
            (
                "1. Resident Community Platform",
                [
                    ("Post Types", "Supports discussions, announcements, polls, and photo albums."),
                    ("Engagement Features", "Supports reactions, comments, poll voting, and post view tracking."),
                    ("Media Posting", "Supports community media upload through dedicated API routes."),
                ],
            ),
            (
                "2. Moderation & Governance",
                [
                    ("Pending Review Queues", "Supports moderation views for pending community content."),
                    ("Approval Controls", "Management roles can approve or reject resident posts."),
                    ("Content Reporting", "Supports reporting inappropriate content and tracking report states."),
                ],
            ),
        ],
    ),
    (
        "X. Analytics, Reporting & Audit",
        [
            (
                "1. Landlord Portfolio Analytics",
                [
                    ("KPI Dashboard", "Tracks earnings, tenants, occupancy, maintenance, renewals, duration, and portfolio value."),
                    ("Trend Visualization", "Supports chart-backed trend views over configurable date ranges."),
                    ("Insight Layer", "Provides KPI insight generation or fallback interpretive guidance."),
                ],
            ),
            (
                "2. Export & Audit Trail",
                [
                    ("CSV Reporting", "Supports exporting landlord portfolio reports as CSV."),
                    ("PDF Reporting", "Supports generating branded PDF portfolio reports."),
                    ("Export History", "Tracks recent export actions for audit visibility."),
                ],
            ),
            (
                "3. Admin Oversight",
                [
                    ("Platform Statistics", "Admin dashboard shows user and registration metrics."),
                    ("User Management", "Admin can inspect platform users through dedicated management interfaces."),
                    ("Registration Governance", "Admin can review and act on landlord registrations."),
                ],
            ),
        ],
    ),
    (
        "XI. Security, Data Integrity & Platform Foundation",
        [
            (
                "1. Data Protection",
                [
                    ("Row-Level Security", "Supabase RLS policies are used across major platform domains."),
                    ("Audit Logging", "Centralized audit event logging captures key request metadata."),
                    ("Signing Security", "Lease workflows use token validation and signature verification utilities."),
                ],
            ),
            (
                "2. Schema & Operational Foundation",
                [
                    ("Role and Policy Controls", "Includes admin-role and community permission matrix support in schema migrations."),
                    ("Workflow-Specific Persistence", "Supports community, lease-signing, registration, and operational data models."),
                    ("Automated Test Coverage", "Includes tests for lease transitions, signature validation, wizard storage, community actions, and selected admin/auth flows."),
                ],
            ),
        ],
    ),
]


def set_normal_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)


def add_bullet(document: Document, label: str, description: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(6)

    run_label = paragraph.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.name = "Times New Roman"
    run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run_label.font.size = Pt(11)

    run_text = paragraph.add_run(description)
    run_text.font.name = "Times New Roman"
    run_text._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run_text.font.size = Pt(11)


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    set_normal_style(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("iReside: System Feature Specification")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(14)

    intro = document.add_paragraph()
    intro.paragraph_format.first_line_indent = Inches(0.5)
    intro.paragraph_format.space_after = Pt(18)
    intro_run = intro.add_run(INTRO)
    intro_run.font.name = "Times New Roman"
    intro_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    intro_run.font.size = Pt(11)

    for section_title, subsections in SECTIONS:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(10)
        heading_run = heading.add_run(section_title)
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        heading_run.font.size = Pt(12)

        for subsection_title, bullets in subsections:
            sub = document.add_paragraph()
            sub.paragraph_format.space_after = Pt(8)
            sub_run = sub.add_run(subsection_title)
            sub_run.bold = True
            sub_run.font.name = "Times New Roman"
            sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            sub_run.font.size = Pt(11)

            for label, description in bullets:
                add_bullet(document, label, description)

            document.add_paragraph().paragraph_format.space_after = Pt(2)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
