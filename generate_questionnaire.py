from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(2)

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

title = doc.add_heading('Feature-based Questionnaire', level=0)
title.runs[0].font.size = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
p.add_run('Directions: ').bold = True
p.add_run('Please pick a rating for each criterion whether: ')
p.add_run('STRONGLY DISAGREE (1), DISAGREE (2), AGREE (3), or STRONGLY AGREE (4).').bold = True

p2 = doc.add_paragraph()
p2.add_run('Panuto: ').bold = True
p2.add_run('Pumili ng antas ng iyong pagsang-ayon para sa bawat pamantayan kung ikaw ay: ')
p2.add_run('LUBOS NA HINDI SUMASANG-AYON (1), HINDI SUMASANG-AYON (2), SUMASANG-AYON (3), o LUBOS NA SUMASANG-AYON (4).').bold = True

doc.add_paragraph()

parts = [
    {
        'title': 'Part I \u2014 Tenant / Bahagi I \u2014 Nangungupahan',
        'statements': [
            ('The login and sign-up process is straightforward and allows users to access the appropriate portal based on their role.',
             'Ang proseso ng login at sign-up ay tuwid at nagpapahintulot na makapasok sa tamang portal base sa papel ng user.'),
            ('The dashboard landing page effectively routes users to the correct interface and displays a welcoming overview based on their assigned role.',
             'Ang dashboard landing page ay tama nagi-redirect ng users at nagpapakita ng buod base sa kanilang papel.'),
            ('The digital lease signing feature provides a clear and user-friendly process for reviewing and electronically signing lease agreements.',
             'Ang digital na lease signing feature ay nagbibigay ng malinaw at user-friendly na proseso para sa pagsusuri at pagpirma ng lease agreements.'),
            ('The lease overview displays all relevant information such as lease terms, start and end dates, renewal eligibility, and current status.',
             'Ang lease overview ay nagpapakita ng lahat ng kaugnay na impormasyon tulad ng lease terms, petsa, renewal eligibility, at status.'),
            ('The document vault provides convenient access to signed master agreements, verified reports, amendments, and other lease-related documents.',
             'Ang document vault ay nagbibigay ng maginhawang access sa mga pirma na master agreement, na-verify na ulat, at iba pang kasulatan.'),
            ('The maintenance request submission feature allows tenants to easily report issues with image uploads and detailed descriptions.',
             'Ang maintenance request feature ay nagpapadali sa mga tenant na mag-report ng isyu na may kasamang imahe at detalyadong paglalarawan.'),
            ('The repair ticket tracking feature provides real-time status updates from pending to completed.',
             'Ang repair ticket tracking ay nagbibigay ng real-time na update sa katayuan mula pending hanggang completed.'),
            ('The real-time messaging platform enables effective communication with the landlord through typing indicators, read receipts, and file sharing.',
             'Ang real-time messaging platform ay nagpapadali ng epektibong komunikasyon sa pamamagitan ng typing indicators, read receipts, at file sharing.'),
            ('The iRis AI assistant provides accurate and helpful responses to questions about building rules, amenities, lease details, and recommendations.',
             'Ang iRis AI assistant ay nagbibigay ng tumpak na sagot sa mga tanong tungkol sa panuntunan ng gusali, amenidad, at detalye ng lease.'),
            ('The financial ledger and payment history display clear breakdowns of rent, water, electricity, and other charges.',
             'Ang financial ledger at payment history ay nagpapakita ng malinaw na detalye ng upa, tubig, kuryente, at iba pang singil.'),
            ('The read-only unit map offers clear visual representation of the building layout for tenant orientation.',
             'Ang read-only na unit map ay nag-aalok ng malinaw na visual na representasyon ng layout ng gusali.'),
            ('The unit transfer request feature provides a straightforward process for requesting transfers to vacant units.',
             'Ang unit transfer request feature ay nagbibigay ng tuwid na proseso para sa paglilipat sa mga bakanteng unit.'),
            ('The move-out request feature allows digital submission of move-out requests with clear status tracking.',
             'Ang move-out request feature ay nagpapahintulot ng digital na pagsusumite ng kahilingan na may malinaw na pagsubaybay.'),
            ('The product tour and onboarding features effectively guide new tenants through the system capabilities.',
             'Ang product tour at onboarding features ay mabisang gumagabay sa mga bagong tenant tungo sa mga kakayahan ng sistema.'),
            ('The notification system effectively prioritizes and delivers relevant alerts about maintenance, payments, and community announcements.',
             'Ang notification system ay mabisang nagbibigay-pangunahin at nagdadala ng mga abiso na may kaugnayan sa maintenance, bayarin, at komunidad.'),
            ('The tenant application tracking with dynamic progress indicators clearly shows application timeline and required actions.',
             'Ang tenant application tracking na may dynamic progress indicators ay malinaw na nagpapakita ng timeline at aksyon.'),
            ('The account management features such as profile viewing and password changes are simple and secure.',
             'Ang mga feature sa pamamahala ng account tulad ng profile at pagpalit ng password ay simple at ligtas.'),
            ('The community hub discussion post feature provides an intuitive way to start and participate in building-wide communication threads.',
             'Ang community hub discussion ay nag-aalok ng madaling paraan upang simulan ang komunikasyon sa buong gusali.'),
            ('The community hub photo album feature with multi-image galleries supports effective visual community updates.',
             'Ang photo album sa community hub na may multi-image gallery ay sumusuporta sa epektibong visual na update sa komunidad.'),
            ('The resident polling mechanism in the community hub provides an effective way to gather and respond to community feedback.',
             'Ang resident poll sa community hub ay nag-aalok ng epektibong paraan upang mangalap ng feedback mula sa komunidad.'),
            ('The community threaded comments and multi-reaction system provide clear and engaging conversations on posts.',
             'Ang threaded comments at multi-reaction system ay nagbibigay ng malinaw at kawili-wiling talakayan sa mga post.'),
            ('The content saving/bookmarking feature in the community hub provides convenient access to important posts and discussions.',
             'Ang feature sa pagse-save/bookmark sa community hub ay nagbibigay ng maginhawang access sa mahahalagang post.'),
        ],
    },
    {
        'title': 'Part II \u2014 Landlord / Bahagi II \u2014 Nagpapaupa',
        'statements': [
            ('The dashboard analytics with AI-driven KPI insights provide meaningful summaries of portfolio performance including earnings, occupancy, and maintenance metrics.',
             'Ang dashboard analytics na may AI-driven na KPI insights ay nagbibigay ng makabuluhang buod ng pagganap ng portfolio kabilang ang kita, occupancy, at maintenance.'),
            ('The simplified and detailed analytics toggle effectively balances comprehensive data with readable presentation for different user preferences.',
             'Ang simplified at detailed analytics toggle ay maayos na nagbabalanse ng komprehensibong data sa madaling basahing presentasyon.'),
            ('The date range filtering with presets and custom selection provides flexible reporting windows.',
             'Ang date range filtering na may mga preset at custom selection ay nagbibigay ng flexible na reporting window.'),
            ('The branded PDF report generation provides professional portfolio performance documents for auditing purposes.',
             'Ang branded na PDF report generation ay nagbibigay ng propesyonal na dokumento ng pagganap ng portfolio para sa auditing.'),
            ('The CSV export and export history features allow convenient data extraction and maintain a chronological log of generated reports.',
             'Ang CSV export at export history features ay nagpapahintulot ng madaling pagkuha ng data at nagpapanatili ng log ng mga ulat.'),
            ('The modular floor planner with drag-and-drop and grid snapping offers precise creation of property layouts.',
             'Ang modular floor planner na may drag-and-drop at grid snap ay nag-aalok ng tiyak na paggawa ng layout ng property.'),
            ('The property and unit management interface allows straightforward creation and editing of property records, unit pricing, deposits, and amenities.',
             'Ang interface sa pamamahala ng property at unit ay nagpapahintulot ng direktang paglikha at pag-edit ng mga rekord, presyo, deposito, at amenidad.'),
            ('The walk-in application processing workflow captures sufficient information for prospective tenant evaluation including documents and checklists.',
             'Ang walk-in application workflow ay nagbibigay ng sapat na impormasyon para sa pagsusuri ng potensyal na tenant.'),
            ('The application status workflow with visual progress indicators clearly tracks application progress from pending to approved.',
             'Ang application status workflow na may visual indicators ay malinaw na nagsubaybay ng pag-unlad mula pending hanggang approved.'),
            ('The lease finalization process creates lease records from approved applications and provisions tenant accounts with clear credential distribution.',
             'Ang proseso ng pag-finalize ng lease ay gumagawa ng lease records mula sa mga approved na aplikasyon at nagpo-provision ng tenant account.'),
            ('The dual-mode lease signing workflow supporting both in-person and remote signing with tenant-first signing order is straightforward and secure.',
             'Ang dual-mode lease signing workflow na sumusuporta sa in-person at remote na pagpirma ay tuwid at ligtas.'),
            ('The maintenance dashboard provides effective tracking of repair tickets with priority-based status from pending through completed.',
             'Ang maintenance dashboard ay nagbibigay ng epektibong pagsubaybay ng repair tickets na may priority-based status.'),
            ('The maintenance ticket rich media galleries with multi-image support provide adequate context for repair assessment.',
             'Ang maintenance ticket rich media galleries na may multi-image support ay nagbibigay ng sapat na konteksto para sa pagsusuri.'),
            ('The invoice management feature allows creation and tracking of invoices with automatic breakdowns of base rent, water, and electricity.',
             'Ang invoice management feature ay nagpapahintulot ng paglikha at pagsubaybay ng mga invoice na may automatic na breakdown.'),
            ('The real-time messaging system with typing indicators, read receipts, and file sharing enables effective communication with tenants.',
             'Ang real-time messaging system na may typing indicators, read receipts, at file sharing ay nagpapadali ng epektibong komunikasyon sa mga tenant.'),
            ('The tenant management features including direct messaging, one-click calling, and financial ledger access provide comprehensive tenant oversight.',
             'Ang tenant management features na may direct messaging, one-click calling, at financial ledger access ay nagbibigay ng komprehensibong pangangasiwa.'),
            ('The renewal eligibility engine dynamically calculates and displays appropriate renewal windows based on lease records.',
             'Ang renewal eligibility engine ay dynamic na nagkukwento at nagpapakita ng angkop na renewal window base sa rekord ng lease.'),
            ('The featured property analytics with total sales, views, and month-over-month growth provides effective property performance tracking.',
             'Ang featured property analytics na may total sales, views, at month-over-month growth ay nagbibigay ng epektibong pagsubaybay.'),
            ('The account management features such as profile viewing and password changes are simple and secure.',
             'Ang mga feature sa pamamahala ng account tulad ng profile at pagpalit ng password ay simple at ligtas.'),
            ('The community hub moderation dashboard with approval queue management and content reporting provides adequate tools for community governance.',
             'Ang community hub moderation dashboard na may approval queue management ay nagbibigay ng sapat na kasangkapan sa governance ng komunidad.'),
            ('The community hub management notices and utility alerts are visually distinct and effectively communicate building-wide announcements.',
             'Ang mga management notices at utility alert sa community hub ay malinaw na nakikita at mabisang nagpapakita ng mga anunsyo sa gusali.'),
            ('The content reporting feature in the community hub enables user-driven moderation for spam and inappropriate behavior.',
             'Ang feature sa pag-uulat ng nilalaman sa community hub ay nagpapahintulot ng tamang moderation na nanggagaling sa user.'),
        ],
    },
    {
        'title': 'Part III \u2014 Super Administrator / Bahagi III \u2014 Super Administrator',
        'statements': [
            ('The live metrics dashboard provides real-time insights into platform health through Total Users, Properties, Leases, and Pending Reviews.',
             'Ang live metrics dashboard ay nagbibigay ng real-time na insights sa kalusugan ng platform na may Total Users, Properties, Leases, at Pending Reviews.'),
            ('The registration pipeline visualization clearly tracks landlord application progress through stages from Pending to Reviewing to Approved.',
             'Ang registration pipeline visualization ay malinaw na nagsubaybay ng pag-unlad ng aplikasyon ng landlord mula Pending hanggang Approved.'),
            ('The user management with role-specific filter combinations and search functionality enables efficient platform-wide user governance.',
             'Ang user management na may role-specific na filter at search ay nagpapadali ng mahusay na pamamahala ng user sa buong platform.'),
            ('The registration review modals with document and photo inspection capabilities provide sufficient information for applicant evaluation decisions.',
             'Ang registration review modal na may kakayahan sa pagsusuri ng dokumento at litrato ay nagbibigay ng sapat na impormasyon para sa mga desisyon.'),
            ('The ability to add internal admin notes and update application status with icon-coded badges enables efficient application processing workflow.',
             'Ang kakayahang magdagdag ng admin notes at mag-update ng application status na may icon-coded badges ay nagpapadali ng maayos na proseso.'),
            ('The secure sign-out mechanism with session termination ensures audit-safe account protection.',
             'Ang ligtas na sign-out mechanism na may session termination ay nagtitiyak ng ligtas na proteksyon ng account.'),
            ('The role-based access control enforcement through Row-Level Security effectively prevents unauthorized cross-role data access.',
             'Ang role-based access control sa pamamagitan ng Row-Level Security ay mabisang pumipigil sa hindi awtorisadong access.'),
            ('The system provides adequate administrative controls and monitoring capabilities to maintain platform governance and operational oversight.',
             'Ang sistema ay nagbibigay ng sapat na mga kontrol at kakayahan sa pagmamanman para mapanatili ang pamamahala ng platform.'),
            ('The account management features such as profile viewing and password changes are simple and secure.',
             'Ang mga feature sa pamamahala ng account tulad ng profile at pagpalit ng password ay simple at ligtas.'),
            ('The system provides adequate guidance and tutorials to support administrative responsibilities.',
             'Ang sistema ay nagbibigay ng sapat na gabay at mga tutorial upang suportahan ang mga tungkuling pang-administratibo.'),
        ],
    },
]

for part in parts:
    h = doc.add_heading(part['title'], level=1)
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1, cols=6, style='Table Grid')

    # Header
    hdr = table.rows[0].cells
    labels = ['No.', 'Statement / Pernisiya', '1', '2', '3', '4']
    for i, txt in enumerate(labels):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        shade_cell(hdr[i], 'D9E2F3')

    for idx, (stmt_en, stmt_fil) in enumerate(part['statements'], 1):
        row = table.add_row()
        cells = row.cells

        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[0].paragraphs[0].add_run(str(idx))
        run.font.size = Pt(9)
        run.bold = True

        p = cells[1].paragraphs[0]
        r_en = p.add_run(stmt_en)
        r_en.font.size = Pt(9)
        r_en.bold = True

        p2 = cells[1].add_paragraph()
        p2.space_before = Pt(1)
        p2.space_after = Pt(1)
        r_fil = p2.add_run(f'({stmt_fil})')
        r_fil.font.size = Pt(9)
        r_fil.italic = True

        for ci in range(2, 6):
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[ci].paragraphs[0].add_run('\u2610')
            run.font.size = Pt(16)

    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(10.5)
        for ci in range(2, 6):
            row.cells[ci].width = Cm(1.5)

    doc.add_paragraph()

output = 'docs/Chapter 3 Docs/FEATURE-BASED-QUESTIONNAIRE/iReside-Feature-Based-Questionnaire.docx'
doc.save(output)
print(f'Saved to {output}')
