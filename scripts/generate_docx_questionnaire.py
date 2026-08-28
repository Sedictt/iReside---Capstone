import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    """Applies background color to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    """Sets inner padding for a table cell (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Applies clean solid black borders to the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_questionnaire_table(doc, actor_title, items):
    """Builds a single actor questionnaire table."""
    # Actor title paragraph
    p_actor = doc.add_paragraph()
    p_actor.paragraph_format.space_before = Pt(12)
    p_actor.paragraph_format.space_after = Pt(4)
    run_actor = p_actor.add_run(actor_title)
    run_actor.font.name = "Times New Roman"
    run_actor.font.size = Pt(11)
    run_actor.font.bold = True
    
    # 5 Columns: Statement (wide), 1, 2, 3, 4
    col_widths = [
        Inches(3.8),   # Statement
        Inches(0.9),   # Strongly Disagree (1)
        Inches(0.85),  # Disagree (2)
        Inches(0.85),  # Agree (3)
        Inches(0.9)    # Strongly Agree (4)
    ]
    
    table = doc.add_table(rows=0, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    # Header Row
    headers = [
        "Statement",
        "STRONGLY\nDISAGREE\n(1)",
        "DISAGREE\n(2)",
        "AGREE (3)",
        "STRONGLY\nAGREE (4)"
    ]
    
    h_row = table.add_row()
    for idx, (cell, text) in enumerate(zip(h_row.cells, headers)):
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        set_cell_shading(cell, "D9E8F5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)
        run.font.bold = True
        
    # Item Rows
    for idx, (eng_text, fil_text) in enumerate(items, start=1):
        row = table.add_row()
        for c_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if c_idx == 0:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                
                # Number & English
                run_num = p.add_run(f"{idx}. {eng_text}\n")
                run_num.font.name = "Times New Roman"
                run_num.font.size = Pt(9.5)
                
                # Filipino translation in italics
                run_fil = p.add_run(f"({fil_text})")
                run_fil.font.name = "Times New Roman"
                run_fil.font.size = Pt(9.0)
                run_fil.font.italic = True
                run_fil.font.color.rgb = RGBColor(50, 50, 50)
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply column widths
    for row in table.rows:
        for idx, w in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = w

def generate_questionnaire_docx():
    out_path = r'c:\Users\JV\Documents\GitHub\iReside\docs\iReside_Feature_Based_Questionnaire.docx'
    doc = Document()
    
    # Configure Standard Letter Portrait
    for s in doc.sections:
        s.orientation = 0  # Portrait
        s.page_width = Inches(8.5)
        s.page_height = Inches(11.0)
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(10)
    run_title = p_title.add_run("3.6.4.2 Feature-based Questionnaire")
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(13)
    run_title.font.bold = True
    run_title.font.italic = True
    
    # Directions
    p_dir = doc.add_paragraph()
    p_dir.paragraph_format.space_before = Pt(0)
    p_dir.paragraph_format.space_after = Pt(4)
    p_dir.paragraph_format.line_spacing = 1.15
    run_d1 = p_dir.add_run("Directions: Please pick a rating for each criterion whether:\n")
    run_d1.font.name = "Times New Roman"
    run_d1.font.size = Pt(10)
    run_d1.font.bold = True
    run_d2 = p_dir.add_run("STRONGLY DISAGREE (1), DISAGREE (2), AGREE (3), or STRONGLY AGREE (4).\n\n")
    run_d2.font.name = "Times New Roman"
    run_d2.font.size = Pt(10)
    
    run_p1 = p_dir.add_run("Panuto: Pumili ng antas ng iyong pagsang-ayon para sa bawat pamantayan kung ikaw ay:\n")
    run_p1.font.name = "Times New Roman"
    run_p1.font.size = Pt(10)
    run_p1.font.bold = True
    run_p2 = p_dir.add_run("LUBOS NA HINDI SUMASANG-AYON (1), HINDI SUMASANG-AYON (2), SUMASANG-AYON (3), o LUBOS NA SUMASANG-AYON (4).")
    run_p2.font.name = "Times New Roman"
    run_p2.font.size = Pt(10)

    # Tenant Items (15 updated items)
    tenant_items = [
        (
            "The private QR code and invite link onboarding process makes it easy and convenient for tenants to join their assigned unit and access lease information.",
            "Pinapadali at ginagawang maginhawa ng pribadong QR code at invite link onboarding process para sa mga tenant ang pagsali sa kanilang nakatalagang unit at pag-access sa impormasyon ng lease."
        ),
        (
            "The digital lease review and electronic signing feature provides a clear and user-friendly way for tenants to execute their lease agreements.",
            "Ang digital lease review at electronic signing feature ay nagbibigay ng malinaw at madaling gamitin na paraan para sa mga tenant na suriin at elektronikong lagdaan ang kanilang kasunduan sa pag-upa."
        ),
        (
            "The lease overview clearly displays all relevant information including monthly rental terms, start and end dates, renewal eligibility, and contractual status.",
            "Ang lease overview ay malinaw na nagpapakita ng lahat ng kaugnay na impormasyon gaya ng halaga ng upa, petsa ng simula at tapos, renewal eligibility, at kasalukuyang status ng kontrata."
        ),
        (
            "The document viewer provides convenient access to download and review signed lease agreements and official property records.",
            "Ang document viewer ay nagbibigay ng maginhawang paraan upang ma-download at masuri ang mga nilagdaang kasunduan sa pag-upa at opisyal na talaan ng property."
        ),
        (
            "The maintenance reporting and tracking features make it easy for tenants to submit repair requests with photos, specify urgency, and monitor progress until completion.",
            "Pinapadali ng maintenance reporting at tracking features para sa mga tenant ang pag-report ng mga sirang gamit na may kasamang litrato, pagtukoy ng antas ng pangangailangan, at pagsubaybay sa proseso hanggang sa maayos."
        ),
        (
            "The direct real-time messaging platform enables effective and transparent communication with the landlord regarding property concerns.",
            "Ang direct real-time messaging platform ay nagbibigay-daan sa epektibo at malinaw na pakikipag-ugnayan sa landlord tungkol sa mga usapin sa property."
        ),
        (
            "The iRis assistant provides helpful, quick guidance on building policies, amenities, payment procedures, and general resident inquiries.",
            "Ang iRis assistant ay nagbibigay ng kapaki-pakinabang at mabilis na gabay tungkol sa mga patakaran ng gusali, pasilidad, proseso ng pagbabayad, at pangkalahatang katanungan ng mga residente."
        ),
        (
            "The financial billing ledger and payment history clearly present itemized breakdowns of monthly base rent, sub-metered water, electricity, and remaining balances.",
            "Ang financial billing ledger at payment history ay malinaw na nagpapakita ng detalyadong breakdown ng buwanang upa, sub-metered na tubig, kuryente, at natitirang balanse."
        ),
        (
            "The GCash and digital payment submission feature provides a fast, convenient, and secure way to upload payment receipts and reference numbers.",
            "Ang GCash at digital payment submission feature ay nagbibigay ng mabilis, maginhawa, at ligtas na paraan upang mag-upload ng resibo at reference number ng bayad."
        ),
        (
            "The advance payment option allows tenants to conveniently submit payments credited toward future billing periods.",
            "Ang advance payment option ay nagbibigay-daan sa mga tenant na maginhawang makapagbayad nang maaga para ma-credit sa mga susunod na buwan ng upa."
        ),
        (
            "The interactive unit map provides an intuitive visual layout of the building floors and unit orientation.",
            "Ang interactive unit map ay nagbibigay ng madaling maintindihang visual layout ng mga palapag ng gusali at lokasyon ng bawat unit."
        ),
        (
            "The lease renewal request feature allows tenants to easily express their intent to extend their tenancy before contract expiration.",
            "Ang lease renewal request feature ay nagbibigay-daan sa mga tenant na madaling magsumite ng kahilingan upang palawigin ang kanilang panunuluyan bago matapos ang kontrata."
        ),
        (
            "The community bulletin board provides an organized channel to view official landlord announcements and building updates.",
            "Ang community bulletin board ay nagbibigay ng maayos na paraan upang makita ang mga opisyal na anunsyo ng landlord at mga update sa gusali."
        ),
        (
            "The notification system effectively delivers timely in-app and email alerts regarding billing statements, payment verifications, and urgent maintenance updates.",
            "Ang notification system ay epektibong naghahatid ng napapanahong in-app at email alerts tungkol sa mga billing statement, kumpirmasyon ng bayad, at mahahalagang maintenance updates."
        ),
        (
            "The first-launch onboarding tour clearly introduces new residents to essential portal tools and navigation upon initial login.",
            "Ang first-launch onboarding tour ay malinaw na nagpapakilala sa mga bagong residente ng mahahalagang feature at paraan ng paggamit ng portal sa kanilang unang pag-login."
        )
    ]
    
    # Landlord Items (15 updated items)
    landlord_items = [
        (
            "The workspace personalization and settings hub allows landlords to easily customize branding, brand theme colors, monogram logos, and business contact details.",
            "Ang workspace personalization at settings hub ay nagbibigay-daan sa mga landlord na madaling i-customize ang branding, kulay ng tema, monogram logo, at detalye ng negosyo."
        ),
        (
            "The unified settings save system with unsaved changes exit protection ensures modifications across all tabs are safely stored without accidental data loss.",
            "Ang unified settings save system na may proteksyon laban sa hindi nai-save na pagbabago ay sumisiguro na ang lahat ng binago sa iba't ibang tab ay ligtas na naitatabi nang hindi nawawala."
        ),
        (
            "The Property and Unit Management interface makes it straightforward to organize buildings, update rental rates, record deposits, and manage amenities.",
            "Pinapadali ng Property at Unit Management interface ang pag-organisa ng mga gusali, pag-update ng presyo ng upa, pagtatala ng deposito, at pamamahala ng mga pasilidad."
        ),
        (
            "The interactive Unit Map provides a clear visual bird's-eye view of unit statuses, occupied spaces, and vacancies across the property.",
            "Ang interactive Unit Map ay nagbibigay ng malinaw na visual layout ng kalagayan ng mga unit, mga okupado, at mga bakanteng espasyo sa buong property."
        ),
        (
            "The Tenant Intake and Private Invite QR Code generator provides a secure and seamless way to onboard verified residents into specific units.",
            "Ang Tenant Intake at Private Invite QR Code generator ay nagbibigay ng ligtas at mabilis na paraan upang ma-onboard ang mga verified na residente sa kanilang partikular na unit."
        ),
        (
            "The Lobby Flyer Studio allows landlords to design professional physical posters with custom photos and export high-resolution print-ready PNG images.",
            "Ang Lobby Flyer Studio ay nagbibigay-daan sa mga landlord na magdisenyo ng propesyonal na poster na may sariling litrato at mag-export ng high-resolution PNG image para sa pag-print."
        ),
        (
            "The Digital Lease drafting and electronic signature workflow enables landlords to generate enforceable contracts and finalize agreements digitally.",
            "Ang digital lease drafting at electronic signature workflow ay nagbibigay-daan sa mga landlord na gumawa ng maayos na kontrata at kumpletuhin ang kasunduan sa pamamagitan ng digital na lagda."
        ),
        (
            "The Sub-meter Utility Reading module simplifies the entry of monthly electricity and water consumption and accurately computes utility bills.",
            "Pinapadali ng Sub-meter Utility Reading module ang pagtatala ng buwanang konsumo sa kuryente at tubig at tumpak na kinakalkula ang mga bayarin."
        ),
        (
            "The Invoice Management module enables batch generation of itemized rent and utility billing statements with automated overdue reminder alerts.",
            "Ang Invoice Management module ay nagbibigay-daan sa sabay-sabay na paggawa ng detalyadong billing ng upa at utility na may automated reminder alerts para sa mga lampas na sa takdang petsa."
        ),
        (
            "The Financial Ledger and Payment Verification feature provides a transparent way to review uploaded tenant receipts and approve payments directly into the accounting records.",
            "Ang Financial Ledger at Payment Verification feature ay nagbibigay ng malinaw na paraan upang suriin ang mga na-upload na resibo ng tenant at aprubahan ang bayad papunta sa accounting records."
        ),
        (
            "The Maintenance Management dashboard simplifies reviewing repair requests, assigning third-party service contractors, and recording final resolution costs into property expenses.",
            "Pinapadali ng Maintenance Management dashboard ang pagsusuri ng mga hiling na kumpuni, pagtatalaga ng mga contractor, at pagtatala ng gastos sa pagkukumpuni bilang property expense."
        ),
        (
            "The Landlord-Tenant Messaging module and Community Announcement publisher facilitate organized, real-time communication and building-wide notices.",
            "Ang Landlord-Tenant Messaging module at Community Announcement publisher ay nagpapadali ng maayos at real-time na pakikipag-ugnayan at pagpapaskil ng mga anunsyo para sa buong gusali."
        ),
        (
            "The iRis Assistant provides valuable portfolio analytics, occupancy summaries, and context-aware smart response drafting to assist daily operations.",
            "Ang iRis Assistant ay nagbibigay ng kapaki-pakinabang na analytics sa portfolio, buod ng occupancy, at maayos na mungkahi ng mensahe upang tumulong sa pang-araw-araw na operasyon."
        ),
        (
            "The Lease Renewal and Move-Out termination workflow provides a structured process to approve contract extensions or release vacant units with inspection records.",
            "Ang Lease Renewal at Move-Out termination workflow ay nagbibigay ng organisadong proseso upang aprubahan ang pagpapalawig ng kontrata o magpalabas ng bakanteng unit kalakip ang talaan ng inspeksyon."
        ),
        (
            "The integrated Documentation Hub and 1-Click Handover Kit provide comprehensive user manuals and technical guides to operate the system independently without dedicated IT support.",
            "Ang integrated Documentation Hub at 1-Click Handover Kit ay nagbibigay ng kumpletong user manual at teknikal na gabay upang mapatakbo ang system nang malaya nang hindi nangangailangan ng hiwalay na IT support."
        )
    ]

    build_questionnaire_table(doc, "Actor: Tenant / Nangungupahan", tenant_items)
    
    # Page break between actors
    doc.add_page_break()
    
    build_questionnaire_table(doc, "Actor: Landlord / Nagpapaupa", landlord_items)
    
    try:
        doc.save(out_path)
        print(f"Successfully generated Questionnaire DOCX at: {out_path}")
    except PermissionError:
        print(f"PermissionError saving {out_path}. Close Word if open.")

if __name__ == "__main__":
    generate_questionnaire_docx()
