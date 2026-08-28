import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    """Applies background color to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    """Sets inner padding for a table cell (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Applies a clean solid border around and inside the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_markdown_test_cases(md_file_path):
    """Parses role sections and test case rows from QA_TEST_CASES.md."""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = []
    current_section = None
    
    for line in content.split('\n'):
        line_clean = line.strip()
        if line_clean.startswith('### '):
            section_name = line_clean.replace('### ', '').strip()
            current_section = {'name': section_name, 'rows': []}
            sections.append(current_section)
        elif line_clean.startswith('| TC_'):
            cols = [c.strip() for c in line_clean.split('|')]
            cols = [c for c in cols if c != '']
            if len(cols) >= 6:
                while len(cols) < 9:
                    cols.append('')
                cols[4] = cols[4].replace('<br>', '\n')
                if current_section:
                    current_section['rows'].append(cols[:9])
                    
    return sections

def create_styled_document(sections, orientation="portrait", out_path="output.docx"):
    """Generates a styled Word document in specified orientation."""
    doc = Document()
    
    # Configure Section Setup
    for s in doc.sections:
        if orientation.lower() == "portrait":
            s.orientation = 0  # Portrait
            s.page_width = Inches(8.5)
            s.page_height = Inches(11.0)
            s.top_margin = Inches(0.4)
            s.bottom_margin = Inches(0.4)
            s.left_margin = Inches(0.35)
            s.right_margin = Inches(0.35)
            # Available table width = 8.5 - 0.7 = 7.8 inches
            col_widths = [
                Inches(1.05),  # Test Case ID
                Inches(0.95),  # Module / Route
                Inches(1.05),  # Test Scenario
                Inches(1.05),  # Preconditions
                Inches(1.30),  # Test Steps
                Inches(1.10),  # Expected Result
                Inches(0.40),  # Actual Result
                Inches(0.45),  # Pass or Failed
                Inches(0.45)   # Remarks
            ]
            font_size_header = Pt(8.5)
            font_size_body = Pt(7.5)
            font_size_title = Pt(10.5)
            font_size_section = Pt(9.5)
        else:
            s.orientation = 1  # Landscape
            s.page_width = Inches(11.0)
            s.page_height = Inches(8.5)
            s.top_margin = Inches(0.4)
            s.bottom_margin = Inches(0.4)
            s.left_margin = Inches(0.4)
            s.right_margin = Inches(0.4)
            # Available table width = 11.0 - 0.8 = 10.2 inches
            col_widths = [
                Inches(1.25),  # Test Case ID
                Inches(1.15),  # Module / Route
                Inches(1.35),  # Test Scenario
                Inches(1.35),  # Preconditions
                Inches(1.70),  # Test Steps
                Inches(1.40),  # Expected Result
                Inches(0.65),  # Actual Result
                Inches(0.65),  # Pass or Failed
                Inches(0.70)   # Remarks
            ]
            font_size_header = Pt(9.5)
            font_size_body = Pt(8.5)
            font_size_title = Pt(12)
            font_size_section = Pt(10)
            
    total_table_width = sum(col_widths)
    
    table = doc.add_table(rows=0, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    # 1. Main Title Banner Row
    title_row = table.add_row()
    title_cell = title_row.cells[0]
    for c in title_row.cells[1:]:
        title_cell.merge(c)
    
    set_cell_shading(title_cell, "D9E8F5")
    set_cell_margins(title_cell, top=120, bottom=120, left=120, right=120)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(
        "Design and Development of iReside: Property and Tenant Services Management System\nfor the Landlords of Valenzuela City"
    )
    run.font.name = "Times New Roman"
    run.font.size = font_size_title
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 2. Table Column Headers
    headers = [
        "Test Case ID",
        "Module / Route",
        "Test Scenario",
        "Preconditions",
        "Test Steps",
        "Expected Result",
        "Actual Result",
        "Pass or Failed",
        "Remarks"
    ]
    
    header_row = table.add_row()
    for idx, (cell, text) in enumerate(zip(header_row.cells, headers)):
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 6, 7, 8] else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = font_size_header
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # 3. Add Role Sections & Test Rows
    for sec in sections:
        # Role Header (Merged across 9 columns)
        p_row = table.add_row()
        p_cell = p_row.cells[0]
        for c in p_row.cells[1:]:
            p_cell.merge(c)
        set_cell_shading(p_cell, "BDD7EE")
        set_cell_margins(p_cell, top=60, bottom=60, left=80, right=80)
        p_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = p_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(sec['name'])
        run.font.name = "Times New Roman"
        run.font.size = font_size_section
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Test Case Rows
        for row_data in sec['rows']:
            d_row = table.add_row()
            for c_idx, (cell, val) in enumerate(zip(d_row.cells, row_data)):
                set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 6, 7, 8] else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run(val)
                run.font.name = "Times New Roman"
                run.font.size = font_size_body
                run.font.color.rgb = RGBColor(0, 0, 0)
                
    # Apply cell widths
    for row in table.rows:
        if len(row.cells) == 1:
            row.cells[0].width = total_table_width
        else:
            for idx, w in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = w

    try:
        doc.save(out_path)
        print(f"Generated {orientation.upper()} at: {out_path}")
    except PermissionError:
        print(f"Note: Could not overwrite {out_path} because it is currently open in Microsoft Word. Close Word to refresh this specific file.")

def generate_all_variants():
    md_path = r'c:\Users\JV\Documents\GitHub\iReside\docs\QA_TEST_CASES.md'
    sections = parse_markdown_test_cases(md_path)
    
    # 1. Letter Portrait (Standard US Letter 8.5" x 11")
    create_styled_document(
        sections,
        orientation="portrait",
        out_path=r'c:\Users\JV\Documents\GitHub\iReside\docs\iReside_QA_Test_Cases_Letter_Portrait.docx'
    )
    
    # 2. Letter Landscape (US Letter 11" x 8.5")
    create_styled_document(
        sections,
        orientation="landscape",
        out_path=r'c:\Users\JV\Documents\GitHub\iReside\docs\iReside_QA_Test_Cases_Letter_Landscape.docx'
    )
    
    # 3. Default unified docx (defaults to Letter Landscape for maximum reading comfort)
    create_styled_document(
        sections,
        orientation="landscape",
        out_path=r'c:\Users\JV\Documents\GitHub\iReside\docs\iReside_QA_Test_Cases.docx'
    )

if __name__ == "__main__":
    generate_all_variants()
